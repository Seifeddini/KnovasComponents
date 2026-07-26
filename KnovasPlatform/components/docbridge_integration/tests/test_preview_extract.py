import sys
from pathlib import Path

SRC = Path(__file__).resolve().parents[1] / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from fixtures.make_msg import build_sample_msg  # noqa: E402


def test_generated_msg_parses_with_extract_msg(tmp_path):
    import extract_msg

    target = tmp_path / "sample.msg"
    build_sample_msg(str(target))

    msg = extract_msg.openMsg(str(target))
    assert msg.subject == "Rückfrage zum Kaufvertrag 2024-001"
    assert msg.sender == "Anna Muster"
    assert msg.isSent is True
    assert msg.date is not None
    assert msg.date.year == 2026 and msg.date.month == 3 and msg.date.day == 15
    assert [r.email for r in msg.recipients] == ["beat.beispiel@example.com"]
    assert "Kaufpreis von EUR 485.000,00" in msg.body


from web_interface.preview import preview_kind  # noqa: E402


def test_preview_kind_recognises_supported_suffixes():
    assert preview_kind("a/b/report.pdf") == "pdf"
    assert preview_kind("Vertrag.DOCX") == "docx"
    assert preview_kind("notiz.txt") == "txt"
    assert preview_kind("mail.msg") == "msg"


def test_preview_kind_rejects_everything_else():
    assert preview_kind("bild.png") is None
    assert preview_kind("archiv.zip") is None
    assert preview_kind("ohne_endung") is None


import pytest  # noqa: E402

from web_interface.preview import (  # noqa: E402
    PreviewFailed,
    PreviewUnsupported,
    extract_markdown,
)


def test_extract_txt(tmp_path):
    target = tmp_path / "notiz.txt"
    target.write_text("Zeile eins.\nZeile zwei mit Umlauten: äöü.\n", encoding="utf-8")

    result = extract_markdown(str(target))
    assert result["kind"] == "txt"
    assert "Zeile eins." in result["markdown"]
    assert "äöü" in result["markdown"]


def test_extract_docx(tmp_path):
    import docx

    target = tmp_path / "vertrag.docx"
    document = docx.Document()
    document.add_heading("Kaufvertrag", 1)
    document.add_paragraph("Der Kaufpreis betraegt EUR 485.000.")
    document.save(str(target))

    result = extract_markdown(str(target))
    assert result["kind"] == "docx"
    assert "# Kaufvertrag" in result["markdown"]
    assert "EUR 485.000" in result["markdown"]
    assert result["meta"]["word_count"] > 0


def test_extract_msg(tmp_path):
    target = tmp_path / "mail.msg"
    build_sample_msg(str(target))

    result = extract_markdown(str(target))
    assert result["kind"] == "msg"
    assert "Kaufpreis von EUR 485.000,00" in result["markdown"]
    assert result["meta"]["title"] == "Rückfrage zum Kaufvertrag 2024-001"
    assert result["meta"]["msg:from"] == "Anna Muster"
    assert result["meta"]["msg:to"] == "beat.beispiel@example.com"


def test_extract_rejects_pdf_and_unknown(tmp_path):
    pdf = tmp_path / "a.pdf"
    pdf.write_bytes(b"%PDF-1.4\n")
    with pytest.raises(PreviewUnsupported):
        extract_markdown(str(pdf))
    with pytest.raises(PreviewUnsupported):
        extract_markdown(str(tmp_path / "bild.png"))


def test_extract_missing_file_raises_preview_failed(tmp_path):
    """A deleted-file race or bad path must surface as PreviewFailed.

    ``knovas_extract.extract`` opens the path with a bare ``open("rb")`` and
    lets ``FileNotFoundError`` (an ``OSError``) escape uncaught. Without a
    broad except clause here, that raw OSError would defeat the typed
    PreviewUnsupported/PreviewFailed contract the Flask route relies on.
    """
    missing = tmp_path / "does-not-exist.txt"
    with pytest.raises(PreviewFailed):
        extract_markdown(str(missing))


def test_extraction_does_not_escape_document_text(tmp_path):
    """Haelt die Annahme fest, auf der der Client-Renderer beruht.

    knovas_extract entfernt *Markup* aus feindlichen Dokumenten, escaped aber
    keinen *Textinhalt*: steht "<script>" woertlich im Fliesstext, kommt es
    unveraendert zurueck. static/js/markdown.js escaped deshalb zuerst und
    formatiert erst danach. Schlaegt dieser Test fehl, weil die Bibliothek
    inzwischen selbst escaped, ist das kein Fehler -- aber der Renderer und
    dieser Kommentar gehoeren dann ueberprueft.
    """
    import docx

    target = tmp_path / "hostile.docx"
    document = docx.Document()
    document.add_paragraph("<script>alert(1)</script>")
    document.save(str(target))

    markdown = extract_markdown(str(target))["markdown"]
    assert "<script>" in markdown
